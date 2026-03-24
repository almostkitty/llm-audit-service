import json
from typing import Any

from docx import Document


def extract_text_from_docx(file_path: str) -> str:
    """Plain text from DOCX (paragraphs only) for the metrics pipeline."""
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_docx_as_json(file_path: str) -> str:
    """DOCX as JSON: paragraph texts (tables and headers are not included)."""
    doc = Document(file_path)
    texts = [p.text for p in doc.paragraphs]
    paragraphs: list[dict[str, Any]] = [
        {"index": i, "text": t} for i, t in enumerate(texts)
    ]
    payload = {
        "source": file_path,
        "paragraphs": paragraphs,
        "plain_text": "\n".join(texts),
    }
    return json.dumps(payload, ensure_ascii=False)
